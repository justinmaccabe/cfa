"""CFA Level II study tracker — the durable, shared record behind the chat tutor.

Run locally:   streamlit run app.py        (uses local SQLite cfa.db, no setup)
Deployed:      set DATABASE_URL secret      (uses Postgres so data persists)

Design note: this app and the chat tutor read/write the SAME database. The app is
where you look; the DB is what the tutor trusts. Neither invents your progress.

Layout: left-sidebar nav (brand + countdown), and a front "Today" page built
around a dated agenda + an exam-runway timeline — deliberately its own identity,
not a portfolio dashboard.
"""
import calendar
import datetime as dt
import glob
import hashlib
import importlib
import os
import re
from io import BytesIO

import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import streamlit as st
import streamlit.components.v1 as components

import curriculum as curr
import db

# Streamlit Cloud keeps a warm process: app.py re-runs on each deploy but sibling
# modules can stay stale in sys.modules, so a newly-added db/curr symbol raises
# AttributeError until reboot. Reload them if a recent symbol is missing.
if not hasattr(db, "READINGS_SEED_OK") or not hasattr(curr, "READINGS"):
    importlib.reload(curr)
    importlib.reload(db)

st.set_page_config(page_title="CFA Level II — Study Tracker",
                   page_icon="◆", layout="wide", initial_sidebar_state="expanded")
db.init_db()

# ----------------------------------------------------------------- palette
PRIMARY = "#335765"   # dark slate blue — headers, rules, structure
TEAL    = "#74A8A4"   # muted teal — positive / strength accent
LBLUE   = "#B6D9E0"   # light blue — soft fills
MIST    = "#DBE2DC"   # pale grey-green — soft background
BROWN   = "#7F543D"   # warm brown — weak / attention accent
INK     = "#262E32"   # near-black body text
MUTE    = "#5B6B72"   # muted label grey
TAPER   = "#5E7E86"   # slate-teal — readable runway segment
SERIF   = "'Times New Roman', Georgia, serif"
POS, NEG = TEAL, BROWN
SIGNAL_COLOR = {"below": BROWN, "at": PRIMARY, "above": TEAL}

st.markdown(f"""
<style>
html, body, [class*="css"], .stApp, [data-testid="stSidebar"], h1, h2, h3, h4, h5,
p, div, span, .stMarkdown, [data-testid="stMetricValue"], [data-testid="stMetricLabel"],
[data-testid="stMetricDelta"], button, input, select, textarea, label {{
    font-family: {SERIF} !important;
}}
span[data-testid="stIconMaterial"], [class*="material-symbols"], .material-icons {{
    font-family: "Material Symbols Rounded" !important;
}}
[data-testid="stMetricValue"], [data-testid="stMetricDelta"],
[data-testid="stDataFrame"], .stDataFrame, table {{
    font-variant-numeric: tabular-nums lining-nums !important;
    font-feature-settings: "tnum" 1, "lnum" 1 !important;
}}
.stApp {{ background:#F5F7F6; }}
.block-container {{ padding-top: 3.6rem; }}
#MainMenu, footer {{ visibility: hidden; }}

/* --- sidebar brand ------------------------------------------------ */
[data-testid="stSidebar"] {{ background:{PRIMARY}; }}
[data-testid="stSidebar"] * {{ color:#EAF0EE !important; }}
.side-mark {{ font-family:{SERIF}; font-style:italic; font-size:2.6rem; font-weight:700;
    color:#FFFFFF !important; line-height:1; letter-spacing:.02em; }}
.side-kick {{ color:{LBLUE} !important; letter-spacing:.34em; text-transform:uppercase;
    font-size:.66rem; margin-top:.2rem; }}
.side-count {{ margin:.9rem 0 .2rem; font-size:2rem; color:#FFFFFF !important; font-weight:700; }}
.side-count small {{ font-size:.8rem; color:{LBLUE} !important; font-weight:400; letter-spacing:.06em; }}
.side-phase {{ display:inline-block; margin-top:.5rem; padding:.18rem .6rem; border-radius:20px;
    background:rgba(255,255,255,.14); color:#FFFFFF !important; font-size:.72rem;
    letter-spacing:.1em; text-transform:uppercase; }}
[data-testid="stSidebar"] hr {{ border-color:rgba(255,255,255,.2); }}
[data-testid="stSidebar"] [role="radiogroup"] label {{ font-size:1.02rem !important; padding:.12rem 0; }}

/* --- headings & rules -------------------------------------------- */
h2, h3 {{ color:{INK}; border-left:3px solid {PRIMARY}; padding-left:.55rem; }}
.page-eyebrow {{ color:{TEAL}; letter-spacing:.34em; text-transform:uppercase; font-size:.72rem; }}
.page-title {{ font-size:2.1rem; font-weight:700; color:{INK}; margin:.1rem 0 .1rem; letter-spacing:.02em; }}
.page-rule {{ border:none; border-top:1px solid rgba(51,87,101,.28); margin:.3rem 0 1.2rem; }}

/* --- the agenda card (signature) --------------------------------- */
.agenda {{ background:#FFFFFF; border:1px solid rgba(51,87,101,.18); border-left:5px solid {PRIMARY};
    border-radius:12px; padding:1.1rem 1.4rem 1.2rem; box-shadow:0 2px 10px rgba(51,87,101,.07); }}
.agenda-date {{ font-size:1.5rem; font-weight:700; color:{PRIMARY}; letter-spacing:.01em; }}
.agenda-tag {{ float:right; font-style:italic; color:{MUTE}; font-size:.95rem; }}
.agenda-item {{ display:flex; align-items:baseline; gap:.6rem; padding:.42rem 0;
    border-top:1px dotted rgba(51,87,101,.16); font-size:1.05rem; color:{INK}; }}
.agenda-item:first-of-type {{ border-top:none; }}
.dot {{ flex:0 0 auto; width:9px; height:9px; border-radius:50%; transform:translateY(1px); }}
.agenda-item .meta {{ margin-left:auto; color:{MUTE}; font-size:.86rem; font-style:italic; white-space:nowrap; }}
.agenda-empty {{ color:{MUTE}; font-style:italic; padding:.4rem 0; }}
.agenda-head {{ display:flex; justify-content:space-between; align-items:baseline; margin:.1rem 0 .35rem; }}

/* --- the exam runway (signature) --------------------------------- */
.runway {{ position:relative; display:flex; height:36px; border-radius:8px; overflow:hidden;
    border:1px solid rgba(51,87,101,.25); margin:.1rem 0 .1rem; }}
.rw-seg {{ display:flex; align-items:center; justify-content:center; min-width:0; }}
.rw-seg span {{ font-size:.68rem; color:#FFFFFF; letter-spacing:.09em; text-transform:uppercase;
    white-space:nowrap; overflow:hidden; text-overflow:ellipsis; padding:0 6px;
    text-shadow:0 1px 1px rgba(0,0,0,.18); }}
.rw-marker {{ position:absolute; top:-3px; bottom:-3px; width:2px; background:{INK}; z-index:3; }}
.rw-marker::after {{ content:'TODAY'; position:absolute; top:-15px; left:50%; transform:translateX(-50%);
    font-size:.58rem; letter-spacing:.12em; color:{INK}; background:#F5F7F6; padding:0 3px; }}
.rw-ends {{ display:flex; justify-content:space-between; color:{MUTE}; font-size:.75rem;
    font-style:italic; margin-top:.5rem; }}
/* runway legend (names + dates, always readable) */
.rw-legend {{ display:flex; flex-wrap:wrap; gap:.35rem 1.1rem; margin-top:.6rem; font-size:.8rem; }}
.rw-leg {{ color:{INK}; white-space:nowrap; }}
.rw-leg b {{ color:{PRIMARY}; }}
.rw-dot {{ display:inline-block; width:10px; height:10px; border-radius:3px; margin-right:.35rem;
    vertical-align:middle; }}
.rw-legd {{ color:{MUTE}; font-style:italic; }}
/* realized vs expected tracker */
.rwp-line {{ font-size:.95rem; color:{INK}; margin-bottom:.35rem; }}
.rwp {{ position:relative; height:12px; border-radius:7px; background:rgba(51,87,101,.12);
    overflow:hidden; }}
.rwp-fill {{ height:100%; background:{TEAL}; border-radius:7px 0 0 7px; }}
.rwp-mark {{ position:absolute; top:-3px; bottom:-3px; width:2px; background:{INK}; }}
.rwp + .runway {{ margin-top:.9rem; }}
/* colour the agenda block: a slate header banner + a slate left edge on the box */
.agenda-head {{ background:{PRIMARY} !important; border-radius:10px 10px 0 0 !important;
    padding:.55rem .95rem !important; margin:.25rem 0 0 !important; align-items:center !important; }}
.agenda-head .agenda-date {{ color:#FFFFFF !important; }}
.agenda-head .agenda-tag {{ color:{LBLUE} !important; float:none !important; }}
[data-testid="stVerticalBlockBorderWrapper"] {{ border-left:4px solid {PRIMARY} !important;
    background:rgba(219,226,220,.28) !important; border-top-left-radius:0 !important;
    border-top-right-radius:0 !important; }}

[data-testid="stMetric"] {{ background:#FFFFFF; border:1px solid rgba(51,87,101,.18);
    border-radius:10px; padding:14px 18px; box-shadow:0 1px 2px rgba(51,87,101,.06); }}
[data-testid="stMetricLabel"] {{ color:{MUTE} !important; text-transform:uppercase;
    letter-spacing:.12em; font-size:.72rem !important; min-height:2.1em; }}
[data-testid="stMetricValue"] {{ font-size:1.6rem !important; color:{INK}; }}

/* --- calendar grid ------------------------------------------------ */
.cal {{ width:100%; border-collapse:collapse; table-layout:fixed; margin-top:.5rem; }}
.cal th {{ color:{MUTE}; font-size:.68rem; text-transform:uppercase; letter-spacing:.12em;
    padding:.3rem 0; text-align:center; font-weight:700; }}
.cal td {{ border:1px solid rgba(51,87,101,.14); vertical-align:top; height:98px;
    padding:4px 5px; background:#FFFFFF; overflow:hidden; }}
.cal td.out {{ background:#EDF1F0; }}
.cal td.past {{ background:#F3F5F4; }}
.cal td.today {{ background:#EAF2F1; box-shadow:inset 0 0 0 2px {PRIMARY}; }}
.cal .dnum {{ font-size:.78rem; color:{INK}; font-weight:700; }}
.cal td.out .dnum {{ color:#AEB7B5; }}
.cal .chip {{ display:block; font-size:.62rem; line-height:1.4; border-radius:4px;
    padding:1px 5px; margin-top:3px; color:#FFFFFF; white-space:nowrap; overflow:hidden;
    text-overflow:ellipsis; }}
.chip.study {{ background:{PRIMARY}; }}
.chip.review {{ background:{BROWN}; }}
.chip.mile {{ background:{TEAL}; }}
.chip.mock {{ background:{TAPER}; }}
.chip.exam {{ background:{BROWN}; font-weight:700; letter-spacing:.04em; }}
.chip.more {{ background:transparent !important; color:{MUTE}; font-style:italic; }}

/* --- curriculum section table --- */
.cur-topic {{ margin:1.15rem 0 .3rem; padding-left:.5rem; border-left:4px solid {PRIMARY};
    font-size:.72rem; letter-spacing:.16em; text-transform:uppercase; color:{PRIMARY}; font-weight:700; }}
.cur-bk {{ display:inline-block; font-size:.62rem; letter-spacing:.05em; color:{MUTE};
    border:1px solid rgba(51,87,101,.3); border-radius:4px; padding:1px 5px; }}
.cbar {{ display:inline-block; width:65%; height:8px; background:rgba(51,87,101,.12);
    border-radius:5px; overflow:hidden; vertical-align:middle; }}
.cbar-fill {{ display:block; height:100%; background:{TEAL}; }}
.cbar-txt {{ font-size:.72rem; color:{MUTE}; margin-left:.45rem; vertical-align:middle; }}

/* --- resource cards --- */
.rescard {{ border:1px solid rgba(51,87,101,.2); border-left:4px solid {PRIMARY};
    border-radius:10px 10px 0 0; background:#FFFFFF; padding:.7rem .9rem .55rem;
    box-shadow:0 1px 2px rgba(51,87,101,.06); }}
.rescard .rc-name {{ color:{INK}; font-weight:700; font-size:1.02rem; }}
.rescard .rc-meta {{ color:{MUTE}; font-size:.8rem; letter-spacing:.04em; }}

/* gentle button polish */
.stButton > button {{ transition: border-color .12s ease, color .12s ease, box-shadow .12s ease; }}
.stButton > button:hover {{ border-color:{PRIMARY}; color:{PRIMARY};
    box-shadow:0 1px 5px rgba(51,87,101,.14); }}
</style>
""", unsafe_allow_html=True)


def _plotly(fig, h=320):
    fig.update_layout(
        template="plotly_white", height=h, margin=dict(l=10, r=10, t=30, b=10),
        paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
        font=dict(family="Times New Roman, Georgia, serif", color=INK),
        legend=dict(orientation="h", y=-0.15))
    return fig


# ----------------------------------------------------------------- data + dates
settings = db.get_settings()
today = dt.date.today()


def _d(key, default):
    return pd.to_datetime(settings.get(key, default)).date()


exam_date = _d("exam_date", "2027-05-18")
pregame_start = _d("pregame_start", "2026-07-20")
materials_date = _d("materials_date", "2026-08-12")
consolidation_start = _d("consolidation_start", "2027-02-03")
mock_start = _d("mock_start", "2027-03-03")
taper_start = _d("taper_start", "2027-05-12")
weekly_target = float(settings.get("weekly_hours_target", 12) or 12)
days_to_exam = (exam_date - today).days

# The runway phases, Pre-game through Taper.
PHASES = [
    ("Pre-game",      pregame_start,       materials_date,      "#A9B2B0"),
    ("Foundation",    materials_date,      consolidation_start, PRIMARY),
    ("Consolidation", consolidation_start, mock_start,          TEAL),
    ("Mocks",         mock_start,          taper_start,         BROWN),
    ("Taper",         taper_start,         exam_date,           TAPER),
]


def current_phase():
    if today < materials_date:
        return "Pre-game"
    for name, s, e, _c in PHASES:
        if s <= today < e:
            return name
    return "Exam week" if today <= exam_date else "Done"


phase = current_phase()


def expected_progress(as_of=None):
    """Where each topic *should* be by `as_of`, and the overall elapsed fraction of
    the content window. Model: lay the non-parallel topics end-to-end across the
    content window (materials_date -> consolidation_start), each occupying a slice
    sized by its CFA weight, in our STUDY_ORDER. A topic's expected % is how far
    'today' has moved through its own slice. Ethics (parallel, order 99) is spread
    linearly across the whole window. Mirrors the xlsx planner's paced dates."""
    as_of = as_of or today
    span = max((consolidation_start - materials_date).days, 1)
    ef = min(max((as_of - materials_date).days, 0), span) / span
    seq = sorted([t for t in curr.TOPICS if curr.STUDY_ORDER.get(t, 50) < 90],
                 key=lambda t: curr.STUDY_ORDER[t])
    w = {t: sum(curr.TOPIC_WEIGHTS[t]) / 2 for t in seq}
    tot = sum(w.values()) or 1
    exp, c = {}, 0.0
    for t in seq:
        s, e = c, c + w[t] / tot
        exp[t] = 0.0 if ef <= s else (1.0 if ef >= e else (ef - s) / (e - s))
        c = e
    for t in curr.TOPICS:            # parallel topics (Ethics) track the window linearly
        exp.setdefault(t, ef)
    return exp, ef


mods = db.get_modules_df()
log = db.get_study_log_df()
mocks_df = db.get_mocks_df()
queue = db.review_queue(today)


# ----------------------------------------------------------------- sidebar
with st.sidebar:
    st.markdown(f"""
    <div class="side-mark">CFA <span style="color:{LBLUE}">II</span></div>
    <div class="side-kick">Level II &middot; May 2027</div>
    <div class="side-count">{max(days_to_exam,0)}<small> days to exam</small></div>
    <span class="side-phase">{phase} phase</span>
    <hr>
    """, unsafe_allow_html=True)
    page = st.radio("Go to", ["Today", "Calendar", "Curriculum", "Notes", "Resources",
                              "Reviews", "Drill Log", "Mocks", "Analytics"],
                    label_visibility="collapsed")
    st.markdown("<hr>", unsafe_allow_html=True)
    st.caption(f"{settings.get('candidate_name','')}  ·  target {weekly_target:.0f} h/wk")


def page_header(eyebrow, title):
    st.markdown(f'<div class="page-eyebrow">{eyebrow}</div>'
                f'<div class="page-title">{title}</div><hr class="page-rule">',
                unsafe_allow_html=True)


# ================================================================= TODAY
def on_agenda_df():
    """The sub-readings that belong on today's agenda — any with a review due, any in
    progress, plus the next few not-started — as one editable slice of the curriculum.
    Editing writes straight back to the sub-module, so whatever you fill in stays filled;
    only the next review's checkbox is blank when the row reappears in 14/45 days."""
    m = db.get_modules_df().sort_values(["study_order", "id"])
    due_ids = list(dict.fromkeys(queue["reading_id"].tolist())) if not queue.empty else []
    active_ids = m[m["status"] == "In Progress"]["id"].tolist()
    nxt_ids = m[m["status"] == "Not Started"]["id"].head(3).tolist()
    order, seen = [], set()
    for i in due_ids + active_ids + nxt_ids:
        if i not in seen:
            order.append(i)
            seen.add(i)
    if not order:
        return m.iloc[0:0]
    df = m[m["id"].isin(order)].copy()
    df["ord"] = df["id"].map({i: n for n, i in enumerate(order)})
    return df.sort_values("ord")


def render_pace():
    """Pace-by-topic bars (actual %) with ◆ markers for required-by-today. Lives on the
    Calendar tab."""
    exp, ef = expected_progress()
    rows = []
    for t in curr.TOPICS:
        sub = mods[mods["topic"] == t]
        pct = (sub["status"] == "Done").mean() * 100 if len(sub) else 0
        lo, hi = curr.TOPIC_WEIGHTS[t]
        rows.append(dict(Topic=t, Done=pct, Expected=exp[t] * 100, Weight=(lo + hi) / 2,
                         Signal=curr.L1_SIGNAL.get(t, "at"), Order=curr.STUDY_ORDER.get(t, 50)))
    prog = pd.DataFrame(rows).sort_values("Order")
    fig = px.bar(prog, x="Done", y="Topic", orientation="h", color="Signal",
                 color_discrete_map=SIGNAL_COLOR, category_orders={"Topic": list(prog["Topic"])},
                 labels={"Done": "% complete"}, hover_data={"Weight": ":.0f"})
    fig.add_trace(go.Scatter(x=prog["Expected"], y=prog["Topic"], mode="markers",
                             marker=dict(symbol="diamond-tall", size=12, color=INK,
                                         line=dict(color="#FFFFFF", width=1)),
                             name="required by today",
                             hovertemplate="should be %{x:.0f}% by today<extra></extra>"))
    fig.update_xaxes(range=[0, 100])
    st.plotly_chart(_plotly(fig, 400), use_container_width=True)
    nmods = mods.groupby("topic")["id"].count().to_dict()
    exp_mods = sum(exp[t] * nmods.get(t, 0) for t in curr.TOPICS)
    done_mods = int((mods["status"] == "Done").sum())
    gap = done_mods - exp_mods
    if ef <= 0:
        verdict = "content phase hasn't started — nothing required yet"
    elif gap >= -0.5:
        verdict = f"on or ahead of pace (+{gap:.1f} sections)"
    else:
        verdict = f":red[{-gap:.1f} sections behind pace]"
    st.caption(f"**◆ = required progress by today** (paced from your plan). "
               f"Completed **{done_mods}** of 42; pace expects ~**{exp_mods:.1f}** → {verdict}.  \n"
               "Bar colour = Level I signal: brown below · slate at · teal above the candidate average.")


def topic_schedule():
    """Per-chapter (topic) expected completion date + on-track flag, from the same
    weight-paced model as the runway. Ethics targets the end of the content window."""
    span = max((consolidation_start - materials_date).days, 1)
    seq = sorted([t for t in curr.TOPICS if curr.STUDY_ORDER.get(t, 50) < 90],
                 key=lambda t: curr.STUDY_ORDER[t])
    w = {t: sum(curr.TOPIC_WEIGHTS[t]) / 2 for t in seq}
    tot = sum(w.values()) or 1
    exp, _ = expected_progress()
    m = db.get_modules_df()          # readings
    out, c = [], 0.0
    for t in seq + ["Ethics"]:
        if t == "Ethics":
            target = consolidation_start
        else:
            c += w[t] / tot
            target = materials_date + dt.timedelta(days=round(c * span))
        st_sub = m[m["topic"] == t]
        done, tt = int((st_sub["status"] == "Done").sum()), len(st_sub)
        dfrac = done / tt if tt else 0
        out.append(dict(topic=t, target=target, done=done, tot=tt,
                        on_track=dfrac >= exp.get(t, 0) - 0.05))
    return out


def current_chapter():
    """The topic (chapter) currently being worked — first active sub-reading, else
    the next not-started. Used as the agenda tag once studying is underway."""
    m = db.get_modules_df().sort_values(["study_order", "id"])
    active = m[m["status"] == "In Progress"]
    if len(active):
        return active.iloc[0]["topic"]
    ns = m[m["status"] == "Not Started"]
    return ns.iloc[0]["topic"] if len(ns) else "All reviewed"


def runway_progress():
    """Realized vs expected progress as *actual* values: realized = share of the 42
    readings actually complete; expected = the weight-paced target for today (the
    real model, not a straight-line assumption)."""
    m = db.get_modules_df()
    n = len(m) or 1
    realized = int((m["status"] == "Done").sum()) / n * 100
    exp, _ = expected_progress()
    cnt = m.groupby("topic")["id"].count().to_dict()
    expected = sum(exp.get(t, 0) * cnt.get(t, 0) for t in curr.TOPICS) / n * 100
    return realized, expected


def runway_html():
    total = max((exam_date - pregame_start).days, 1)
    segs = "".join(
        f'<div class="rw-seg" style="width:{max((e - s).days, 0) / total * 100:.2f}%;'
        f'background:{color}"></div>' for name, s, e, color in PHASES)
    frac = min(max((today - pregame_start).days / total, 0), 1)
    marker = f'<div class="rw-marker" style="left:{frac * 100:.2f}%"></div>'
    legend = "".join(
        f"<span class='rw-leg'><span class='rw-dot' style='background:{color}'></span>"
        f"<b>{name}</b> <span class='rw-legd'>{s:%b %-d} – {e:%b %-d, %Y}</span></span>"
        for name, s, e, color in PHASES)
    r, e = runway_progress()
    delta = r - e
    if e < 0.5:
        emoji, msg = "⚪", "the clock hasn't started — pre-game"
    elif delta >= 0:
        emoji, msg = "🟢", f"ahead of pace (+{delta:.0f} pts)"
    elif delta >= -5:
        emoji, msg = "🟡", f"a touch behind ({delta:.0f} pts)"
    else:
        emoji, msg = "🔴", f"behind pace ({delta:.0f} pts)"
    tracker = (f"<div class='rwp-line'>🎯 Realized <b>{r:.0f}%</b> &nbsp;·&nbsp; "
               f"◆ Expected <b>{e:.0f}%</b> &nbsp;&nbsp; {emoji} {msg}</div>"
               f"<div class='rwp'><div class='rwp-fill' style='width:{r:.1f}%'></div>"
               f"<div class='rwp-mark' style='left:{e:.1f}%'></div></div>")
    return (tracker + f'<div class="runway">{segs}{marker}</div>'
            f'<div class="rw-legend">{legend}</div>')


def page_today():
    page_header("Daily brief", f"{today:%A}, {today:%B %-d}")

    started_any = (db.get_submodules_df()["status"] != "Not Started").any()
    tag = current_chapter() if (started_any or today >= materials_date) else "Pre-game phase"
    st.markdown(f"<div class='agenda-head'><span class='agenda-date'>📋 On the agenda</span>"
                f"<span class='agenda-tag'>{tag}</span></div>", unsafe_allow_html=True)
    with st.container(border=True):
        cl = db.get_checklist()
        todo = cl[~cl["done"].astype(bool)]
        for _, t in todo.iterrows():
            cc = st.columns([0.05, 0.95], vertical_alignment="center")
            if cc[0].checkbox("x", key=f"cl_{t['key']}", label_visibility="collapsed"):
                db.set_checklist_done(t["key"], True, today)
                st.rerun()
            cc[1].markdown(f"{t['label']}  \n:gray[{t['note']}]")

        ag = on_agenda_df()
        if ag.empty:
            if todo.empty:
                st.caption("Nothing on the agenda — you're all caught up. 🎯")
        else:
            if not todo.empty:
                st.markdown("<hr style='border:none;border-top:1px dotted rgba(51,87,101,.2);"
                            "margin:.5rem 0'>", unsafe_allow_html=True)
            st.caption("Click a reading to open its modules · Key Concepts · Module Quiz. "
                       "The box on the right clears it from here once you're done.")
            for _, r in ag.iterrows():
                due = queue[queue["reading_id"] == r["id"]] if not queue.empty else None
                has_due = due is not None and len(due)
                why = (f"review {due.iloc[0]['review']} due" if has_due
                       else ("in progress" if r["status"] == "In Progress" else "up next"))
                cc = st.columns([6.2, 2.6, 0.7], vertical_alignment="center")
                if cc[0].button(r["name"], key=f"ag_open_{r['id']}", width="stretch"):
                    section_dialog(int(r["id"]), r["name"], r["topic"])
                cc[1].markdown(f"<span style='color:{PRIMARY};font-style:italic'>"
                               f"{r['topic']} · {why}</span>", unsafe_allow_html=True)
                if cc[2].checkbox("done", key=f"ag_clr_{r['id']}", label_visibility="collapsed"):
                    if has_due:
                        db.update_module(int(r["id"]),
                                         **{f"{due.iloc[0]['review'].lower()}_done": True})
                    else:
                        for sid in db.submodules_for_section(int(r["id"]))["id"]:
                            db.update_submodule(int(sid), status="Practice Complete")
                    st.rerun()

    st.write("")
    st.markdown("##### Exam runway")
    st.markdown(runway_html(), unsafe_allow_html=True)

    st.write("")
    done = int((mods["status"] == "Done").sum())
    in_prog = int((mods["status"] == "In Progress").sum())
    week_start = today - dt.timedelta(days=today.weekday())
    hrs_week = 0.0
    if not log.empty:
        lg = log.copy(); lg["date"] = pd.to_datetime(lg["date"]).dt.date
        hrs_week = lg[lg["date"] >= week_start]["minutes"].sum() / 60
    overdue = 0 if queue.empty else int((queue["days_overdue"] > 0).sum())

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Readings complete", f"{done} / 42", f"{done/42*100:.0f}%")
    c2.metric("In progress", in_prog)
    c3.metric("Hours this week", f"{hrs_week:.1f}", f"target {weekly_target:.0f}")
    c4.metric("Reviews due", 0 if queue.empty else len(queue),
              f"{overdue} overdue" if overdue else "on track", delta_color="inverse")

    st.markdown("<hr class='page-rule'>", unsafe_allow_html=True)
    st.subheader("Schedule")
    cA, cB = st.columns([1, 1])
    with cA:
        st.markdown("**Due this week**")
        ev = build_events()
        wk = sorted((d, k, t) for d, its in ev.items() for k, t in its
                    if today <= d <= today + dt.timedelta(days=7))
        q7 = db.review_queue(today + dt.timedelta(days=7))
        shown = False
        for d, k, t in wk:
            st.markdown(f":gray[{d:%a %b %-d}] · {t}")
            shown = True
        if not q7.empty:
            for _, r in q7.iterrows():
                st.markdown(f":gray[{r['due']:%a %b %-d}] · Review {r['review']} · {r['module']}")
                shown = True
        if not shown:
            st.caption("Nothing due in the next 7 days.")
    with cB:
        st.markdown("**Chapter pace — expected completion**")
        for row in topic_schedule():
            badge = ":green[on track]" if row["on_track"] else ":red[behind]"
            st.markdown(f"**{row['topic']}** · target {row['target']:%b %-d, %Y} · {badge}")
            st.markdown(_bar(row["done"] / row["tot"] * 100 if row["tot"] else 0,
                             row["done"], row["tot"]), unsafe_allow_html=True)


# ================================================================= CURRICULUM
@st.dialog("Reading", width="large")
def section_dialog(sec_id, sec_name, topic):
    st.markdown(f"<span class='page-eyebrow'>{topic}</span><br><b>{sec_name}</b>",
                unsafe_allow_html=True)
    subs = db.submodules_for_section(sec_id)
    show = subs[["id", "code", "name", "status", "confidence", "notes"]].copy()
    fx_col = subs["formulas"].fillna("").apply(
        lambda s: f"{len([ln for ln in s.splitlines() if ln.strip()])} saved" if s.strip() else "—")
    show["formulas"] = fx_col
    edited = st.data_editor(
        show, width="stretch", hide_index=True, key=f"subed_{sec_id}",
        column_config={
            "id": None,
            "code": st.column_config.TextColumn("#", disabled=True, width="small"),
            "name": st.column_config.TextColumn("Item", disabled=True, width="large"),
            "status": st.column_config.SelectboxColumn(
                "Status", options=curr.ITEM_STATUS_OPTIONS, width="medium"),
            "confidence": st.column_config.NumberColumn(
                "Perf %", min_value=0, max_value=100, step=1, format="%d%%", width="small"),
            "notes": st.column_config.TextColumn("Notes", width="medium"),
            "formulas": st.column_config.TextColumn(
                "Formulas", disabled=True, width="medium",
                help="How many formulas are saved — read/edit the full text in the Formulas box below."),
        })
    st.caption("When every item reaches **Practice Complete**, the reading is done and its "
               "+3 / +14 / +45-day reviews arm (Reviews tab). **Module Quiz** performance flows "
               "to the Drill Log → Analytics.")
    if st.button("Save", type="primary", key=f"savesub_{sec_id}"):
        orig, n = show.set_index("id"), 0
        for _, r in edited.iterrows():
            o, ch = orig.loc[r["id"]], {}
            for col in ["status", "confidence", "notes"]:
                new, old = r[col], o[col]
                if pd.isna(new) and pd.isna(old):
                    continue
                if new != old:
                    ch[col] = None if pd.isna(new) else new
            if ch:
                db.update_submodule(int(r["id"]), **ch)
                n += 1
                if r["name"] == "Module Quiz" and "confidence" in ch and pd.notna(r["confidence"]):
                    db.log_study(today, topic, "Practice", num_q=100,
                                 num_correct=int(r["confidence"]), source="Module Quiz",
                                 notes=sec_name)
        st.success(f"Saved {n} change(s).")
        st.rerun()

    # --- Formulas: multi-line, one per line, per item -------------------------
    st.markdown("<hr class='page-rule'>", unsafe_allow_html=True)
    st.markdown("##### 📐 Formulas")
    opts = {(f"{r['code']}  {r['name']}" if r["code"] else r["name"]): int(r["id"])
            for _, r in subs.iterrows()}
    label = st.selectbox("Item", list(opts.keys()), key=f"fx_sel_{sec_id}")
    iid = opts[label]
    cur = subs.set_index("id").loc[iid, "formulas"]
    cur = "" if pd.isna(cur) else cur
    new_fx = st.text_area("One formula per line", cur, height=180, key=f"fx_txt_{sec_id}_{iid}",
                          placeholder="R² = (Total Var − Unexplained Var) / Total Var\n"
                                       "Adjusted R² = 1 − [(n−1)/(n−k−1)]·(1−R²)\nAIC = …")
    if st.button("Save formulas", key=f"fx_save_{sec_id}"):
        db.update_submodule(iid, formulas=new_fx.strip() or None)
        st.success("Formulas saved.")
        st.rerun()


def _bar(fill, done, tot):
    """Bar fill % and the done/tot label are decoupled: on Curriculum the fill shows
    reading progress while the number stays practice-complete."""
    return (f"<div class='cbar'><div class='cbar-fill' style='width:{fill:.0f}%'></div></div>"
            f"<span class='cbar-txt'>{done}/{tot}</span>")


def page_curriculum():
    n_ch, n_read, n_items = len(curr.TOPICS), len(curr.READINGS), len(curr.ITEMS)
    page_header("The map",
                f"Curriculum <i style='font-weight:400;font-size:1.35rem;color:{MUTE}'>"
                f"· {n_ch} Chapters · {n_read} Readings · {n_items} Items</i>")
    st.caption("Each row is a reading. **Open** it to check off its modules, Key Concepts and "
               "Module Quiz — reading progress, the calendar and the pace bars all roll up from there.")
    subs_all = db.get_submodules_df()
    cur_topic = None
    for _, s in mods.sort_values(["study_order", "id"]).iterrows():
        if s["topic"] != cur_topic:
            cur_topic = s["topic"]
            sig = curr.L1_SIGNAL.get(cur_topic, "at")
            st.markdown(f"<div class='cur-topic' style='border-color:{SIGNAL_COLOR[sig]}'>"
                        f"{cur_topic}</div>", unsafe_allow_html=True)
        sub = subs_all[subs_all["section_id"] == s["id"]]
        mod_items = sub[sub["code"].fillna("").astype(str).str.strip() != ""]  # 1.1, 1.2 …
        read_done = int(mod_items["status"].isin(curr.READ_DONE_STATES).sum())
        fill = read_done / len(mod_items) * 100 if len(mod_items) else 0        # bar = reading
        done, tot = int(sub["status"].isin(curr.ITEM_COMPLETE).sum()), len(sub)  # number = practice
        c = st.columns([0.7, 6, 2, 1.3], vertical_alignment="center")
        c[0].markdown(f"<span class='cur-bk'>Bk {int(s['book'])}</span>", unsafe_allow_html=True)
        c[1].markdown(f"**{s['name']}**")
        c[2].markdown(_bar(fill, done, tot), unsafe_allow_html=True)
        if c[3].button("Open ▸", key=f"open_{s['id']}", width="stretch"):
            section_dialog(int(s["id"]), s["name"], s["topic"])


# ================================================================= REVIEWS
def page_reviews():
    page_header("Spaced retrieval", "Review queue")
    st.caption("Reviews are **retrieval practice**, not re-reading. Mark done as you clear them.")
    with st.expander("How to run a review — what the evidence says"):
        st.markdown(
            "1. **Closed-book first.** Do a set of practice questions / blue-box problems on "
            "this reading *from memory* — no notes open. This is the part that builds durable "
            "recall (Roediger & Karpicke; Dunlosky rates practice testing *high-utility*).\n"
            "2. **Then check & score.** Grade yourself, log the % in the Drill Log.\n"
            "3. **Repair the gaps only.** For what you missed, re-study *just* those points, "
            "then re-test them — don't re-read the whole chapter (low-utility).\n"
            "4. **Space it out.** The +3 / +14 / +45-day gaps are deliberate — the struggle to "
            "recall after forgetting is what strengthens memory (Cepeda spacing effect).\n\n"
            "*Reviewing your notes or re-reading feels productive but is among the least "
            "effective methods — keep review = testing yourself.*")
    full_q = db.review_queue(today + dt.timedelta(days=7))
    if full_q.empty:
        st.write("Nothing scheduled yet — finish a reading (Curriculum → open one, complete its "
                 "items) to start the clock.")
        return
    for _, r in full_q.iterrows():
        tag = (f":red[{r['days_overdue']}d overdue]" if r["days_overdue"] > 0
               else (":orange[due today]" if r["days_overdue"] == 0
                     else f":gray[in {-r['days_overdue']}d]"))
        cols = st.columns([5, 1.4, 1.4, 1.2], vertical_alignment="center")
        cols[0].write(f"**{r['module']}**  \n:gray[{r['topic']} · {r['review']}]")
        cols[1].write(f"{r['due']:%b %d}")
        cols[2].write(tag)
        if cols[3].button("Done", key=f"rev_{r['reading_id']}_{r['review']}"):
            db.update_module(int(r["reading_id"]), **{f"{r['review'].lower()}_done": True})
            db.log_study(today, r["topic"], "Review", source="review-queue",
                         notes=f"{r['module']} {r['review']} complete")
            st.rerun()


# ================================================================= DRILL LOG
def page_drill():
    page_header("Practice", "Drill log")
    with st.form("drill_form", clear_on_submit=True):
        c = st.columns(4)
        d_date = c[0].date_input("Date", today)
        d_topic = c[1].selectbox("Topic", curr.TOPICS)
        d_act = c[2].selectbox("Activity", ["Practice", "Read", "Review", "Class", "Other"])
        d_min = c[3].number_input("Minutes", min_value=0, value=30, step=5)
        c2 = st.columns(4)
        d_predict = c2[0].number_input("Predicted % (gut feel)", 0, 100, 60, 5,
                                       help="Your guess BEFORE grading — powers the calibration chart")
        d_nq = c2[1].number_input("Questions", min_value=0, value=0, step=1)
        d_nc = c2[2].number_input("Correct", min_value=0, value=0, step=1)
        d_src = c2[3].selectbox("Source", ["Claude", "CFAI QBank", "Schweser", "Mock", "Other"])
        d_notes = st.text_input("Notes", "")
        if st.form_submit_button("Log it", type="primary"):
            db.log_study(d_date, d_topic, d_act, minutes=d_min,
                         num_q=d_nq or None, num_correct=d_nc or None,
                         predicted=d_predict, source=d_src, notes=d_notes or None)
            st.success("Logged.")
            st.rerun()

    st.subheader("Recent sessions")
    if log.empty:
        st.write("No sessions logged yet.")
        return
    disp = log.copy()
    disp["date"] = pd.to_datetime(disp["date"]).dt.date
    disp["acc %"] = (disp["num_correct"] / disp["num_q"] * 100).round(0)
    st.dataframe(disp[["date", "topic", "activity", "minutes", "num_q", "num_correct",
                       "acc %", "predicted", "source", "notes"]].sort_values("date", ascending=False),
                 width="stretch", hide_index=True, height=300)


# ================================================================= MOCKS
def page_mocks():
    page_header("Under exam conditions", "Mock exams")
    with st.form("mock_form", clear_on_submit=True):
        c = st.columns(4)
        m_date = c[0].date_input("Date", today, key="m_date")
        m_src = c[1].text_input("Source", "CFAI Mock A")
        m_score = c[2].number_input("Score %", 0.0, 100.0, 65.0, 0.5)
        m_time = c[3].number_input("Minutes used", 0, 300, 132)
        m_weak = st.text_input("Weakest topics", "")
        m_action = st.text_input("Action items", "")
        if st.form_submit_button("Add mock", type="primary"):
            db.add_mock(m_date, m_src, m_score, m_time, m_weak or None, m_action or None)
            st.success("Mock added.")
            st.rerun()

    if mocks_df.empty:
        return
    md = mocks_df.copy()
    md["date"] = pd.to_datetime(md["date"]).dt.date
    fig = px.line(md, x="date", y="score_pct", markers=True,
                  labels={"score_pct": "Score %", "date": ""})
    fig.update_traces(line_color=PRIMARY)
    fig.add_hline(y=70, line_dash="dot", line_color=TEAL, annotation_text="70% comfort line")
    st.plotly_chart(_plotly(fig), use_container_width=True)
    st.dataframe(md[["date", "source", "score_pct", "minutes_used",
                     "weak_topics", "action_items", "reviewed"]],
                 width="stretch", hide_index=True)


# ================================================================= ANALYTICS
def page_analytics():
    page_header("Where the points are", "Analytics")
    graded = log[log["num_q"].notna() & (log["num_q"] > 0)].copy() if not log.empty else pd.DataFrame()

    st.subheader("Topic mastery")
    if graded.empty:
        st.write("Log some graded practice sets and this fills in — accuracy by topic, "
                 "plus whether your confidence matches reality.")
    else:
        graded["acc"] = graded["num_correct"] / graded["num_q"] * 100
        by_t = (graded.groupby("topic")
                .apply(lambda g: pd.Series({
                    "accuracy": (g["num_correct"].sum() / g["num_q"].sum() * 100),
                    "questions": int(g["num_q"].sum())}))
                .reset_index())
        by_t["signal"] = by_t["topic"].map(curr.L1_SIGNAL)
        fig = px.bar(by_t.sort_values("accuracy"), x="accuracy", y="topic", orientation="h",
                     color="signal", color_discrete_map=SIGNAL_COLOR,
                     hover_data={"questions": True}, labels={"accuracy": "Accuracy %"})
        fig.add_vline(x=70, line_dash="dot", line_color=PRIMARY)
        fig.update_xaxes(range=[0, 100])
        st.plotly_chart(_plotly(fig, 380), use_container_width=True)

        st.subheader("Confidence calibration")
        st.caption("Above the line = you underestimate yourself; below = overconfident. "
                   "For a re-taker, the overconfident topics are where points quietly leak.")
        cal = graded[graded["predicted"].notna()].copy()
        if cal.empty:
            st.write("Add a *predicted %* when you log a set to unlock this.")
        else:
            fig = go.Figure()
            fig.add_trace(go.Scatter(x=[0, 100], y=[0, 100], mode="lines",
                                     line=dict(dash="dot", color=TEAL), name="perfect calibration"))
            fig.add_trace(go.Scatter(x=cal["predicted"], y=cal["acc"], mode="markers",
                                     marker=dict(size=11, color=PRIMARY), text=cal["topic"],
                                     name="your sets",
                                     hovertemplate="%{text}<br>predicted %{x}%<br>actual %{y:.0f}%"))
            fig.update_layout(xaxis_title="Predicted %", yaxis_title="Actual %")
            fig.update_xaxes(range=[0, 100]); fig.update_yaxes(range=[0, 100])
            st.plotly_chart(_plotly(fig), use_container_width=True)

    st.subheader("Study hours by week")
    if log.empty:
        st.write("No hours logged yet.")
        return
    h = log.copy()
    h["date"] = pd.to_datetime(h["date"])
    h["week"] = h["date"].dt.to_period("W").dt.start_time
    wk = h.groupby("week")["minutes"].sum().reset_index()
    wk["hours"] = wk["minutes"] / 60
    fig = px.bar(wk, x="week", y="hours", labels={"hours": "Hours", "week": ""})
    fig.update_traces(marker_color=PRIMARY)
    fig.add_hline(y=weekly_target, line_dash="dot", line_color=TEAL,
                  annotation_text=f"{weekly_target:.0f} h target")
    st.plotly_chart(_plotly(fig), use_container_width=True)


# ================================================================= CALENDAR
ABBR = {"Quantitative Methods": "QM", "Economics": "ECON",
        "Financial Statement Analysis": "FSA", "Corporate Issuers": "CI",
        "Equity Investments": "EQ", "Fixed Income": "FI", "Derivatives": "DER",
        "Alternative Investments": "ALT", "Portfolio Management": "PM", "Ethics": "ETH"}


def build_events():
    """date -> [(kind, text)] for everything scheduled between materials and exam:
    weight-paced module study-days (same model as the runway/pace bars), spaced
    reviews off completions, phase milestones, and suggested mock slots."""
    ev = {}

    def add(d, kind, text):
        ev.setdefault(d, []).append((kind, text))

    add(materials_date, "mile", "Materials unlock")
    add(dt.date(2026, 10, 14), "mile", "Early-reg deadline")
    add(consolidation_start, "mile", "Consolidation buffer")
    add(mock_start, "mile", "Mocks begin")
    add(taper_start, "mile", "Taper begins")
    add(exam_date, "exam", "EXAM DAY")

    span = max((consolidation_start - materials_date).days, 1)
    n_by_topic = mods.groupby("topic")["id"].count().to_dict()
    seq = mods[mods["study_order"] < 90].sort_values(["study_order", "id"])
    wts = [(sum(curr.TOPIC_WEIGHTS[r["topic"]]) / 2) / n_by_topic[r["topic"]]
           for _, r in seq.iterrows()]
    total = sum(wts) or 1
    cum = 0.0
    for (_, r), wt in zip(seq.iterrows(), wts):
        d = materials_date + dt.timedelta(days=round(cum / total * span))
        add(d, "study", f"{ABBR[r['topic']]} · {r['name']}")
        cum += wt
    for i, (_, r) in enumerate(mods[mods["study_order"] >= 90].sort_values("id").iterrows()):
        frac = [0.30, 0.60, 0.90][i] if i < 3 else 0.9
        add(materials_date + dt.timedelta(days=round(frac * span)),
            "study", f"ETH · {r['name']}")

    # (spaced reviews are tracked at the sub-module level on the Reviews tab, not
    # here — the calendar deliberately stays at the coarser section level.)
    d = mock_start
    while d < taper_start:
        add(d, "mock", "Suggested mock")
        d += dt.timedelta(days=14)
    return ev


def month_html(year, month, ev):
    cal = calendar.Calendar(firstweekday=0)
    out = ['<table class="cal"><thead><tr>']
    out += [f"<th>{w}</th>" for w in ["Mon", "Tue", "Wed", "Thu", "Fri", "Sat", "Sun"]]
    out.append("</tr></thead><tbody>")
    for week in cal.monthdatescalendar(year, month):
        out.append("<tr>")
        for d in week:
            cls = []
            if d.month != month:
                cls.append("out")
            elif d == today:
                cls.append("today")
            elif d < today:
                cls.append("past")
            items = ev.get(d, [])
            chips = "".join(f'<span class="chip {k}">{t}</span>' for k, t in items[:3])
            if len(items) > 3:
                chips += f'<span class="chip more">+{len(items) - 3} more</span>'
            out.append(f'<td class="{" ".join(cls)}"><span class="dnum">{d.day}</span>{chips}</td>')
        out.append("</tr>")
    out.append("</tbody></table>")
    return "".join(out)


def page_calendar():
    page_header("The plan, day by day", "Calendar")
    st.subheader("Pace by topic")
    render_pace()
    st.markdown("<hr class='page-rule'>", unsafe_allow_html=True)
    st.subheader("Month view")
    ev = build_events()
    months, y, m = [], today.year, today.month
    while (y, m) <= (exam_date.year, exam_date.month):
        months.append((y, m))
        y, m = (y + 1, 1) if m == 12 else (y, m + 1)
    labels = [dt.date(yy, mm, 1).strftime("%B %Y") for (yy, mm) in months]
    sel = st.selectbox("Month", range(len(months)), format_func=lambda k: labels[k])
    st.markdown(
        f'<div style="margin:.1rem 0;font-size:.75rem">'
        f'<span class="chip study" style="display:inline-block;color:#FFFFFF">study</span> '
        f'<span class="chip review" style="display:inline-block;color:#FFFFFF">review</span> '
        f'<span class="chip mock" style="display:inline-block;color:#FFFFFF">mock</span> '
        f'<span class="chip mile" style="display:inline-block;color:#FFFFFF">milestone</span></div>',
        unsafe_allow_html=True)
    yy, mm = months[sel]
    st.markdown(month_html(yy, mm, ev), unsafe_allow_html=True)

    rows = sorted((d, k, t) for d, its in ev.items() for (k, t) in its
                  if d.year == yy and d.month == mm)
    if rows:
        st.markdown("##### This month, in order")
        for d, k, t in rows:
            st.markdown(f"- **{d:%a %b %-d}** · :gray[{k}] — {t}")


# ================================================================= access gate
def require_unlock():
    """Password gate for materials + notes. The SHA-256 hash lives in the DB, so the
    password is set once and syncs across devices; no plaintext is ever stored."""
    h = settings.get("resources_pw_sha256") or ""
    if not h:
        st.info("Set a password to protect your materials — only you'll know it. "
                "It syncs across your devices and can be changed later.")
        p = st.text_input("Create a password", type="password", key="pw_set")
        if st.button("Set password", type="primary") and p:
            db.set_setting("resources_pw_sha256", hashlib.sha256(p.encode()).hexdigest())
            st.session_state.unlocked = True
            st.rerun()
        return False
    if st.session_state.get("unlocked"):
        return True
    p = st.text_input("Password", type="password", key="pw_in")
    if st.button("Unlock", type="primary"):
        if hashlib.sha256(p.encode()).hexdigest() == h:
            st.session_state.unlocked = True
            st.rerun()
        else:
            st.error("Wrong password.")
    return False


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


# ================================================================= NOTES (living Word doc)
def _notes_filename(now=None):
    """Filenames are always stamped with the save time, whatever the upload was called:
    CFA-L2-Notes-YYYY-MM-DD-HH-MM.docx."""
    now = now or dt.datetime.now()
    return f"CFA-L2-Notes-{now:%Y-%m-%d-%H-%M}.docx"


def _starter_docx() -> bytes:
    from docx import Document
    from docx.shared import RGBColor
    accent = RGBColor(0x5E, 0x7E, 0x86)   # muted slate-teal for sub-reading bullets
    doc = Document()
    doc.add_heading("CFA Level II — Study Notes", 0)
    doc.add_paragraph("Living notes. Each topic is a chapter — open Word's Navigation Pane "
                      "(View ▸ Navigation Pane) to jump straight to any topic. Each module "
                      "lists its Schweser sub-readings; write your notes beneath them.")
    subs = db.get_submodules_df()
    last_topic = None
    for _, s in mods.sort_values(["study_order", "id"]).iterrows():
        if s["topic"] != last_topic:
            if last_topic is not None:
                doc.add_page_break()
            doc.add_heading(s["topic"], level=1)        # chapter -> Navigation Pane
            last_topic = s["topic"]
        doc.add_heading(s["name"], level=2)             # module
        for _, sm in subs[subs["section_id"] == s["id"]].iterrows():
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"{sm['code']}   {sm['name']}")
            run.font.color.rgb = accent
            run.italic = True
        doc.add_paragraph("")                           # room to write
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def page_notes():
    page_header("Cross-device workspace", "Notes")
    if not require_unlock():
        return
    fname, data = db.get_notes()
    st.caption("Your living notes doc. **Download** it, edit in Word, then **upload** to "
               "override — every session builds on the last, on any computer. Each save is "
               "stamped and kept in the history below as a safety net.")
    c1, c2 = st.columns([1, 1])
    if data:
        c1.download_button(f"⬇  Download current · {fname}", data=data, file_name=fname,
                           mime=DOCX_MIME, width="stretch")
    else:
        if c1.button("Generate starter doc", type="primary", width="stretch"):
            now = dt.datetime.now()
            db.add_notes(_notes_filename(now), _starter_docx(), now.isoformat(timespec="minutes"))
            st.success("Starter doc created — download it below.")
            st.rerun()
    up = st.file_uploader("Upload a new version (overrides current)", type=["docx"])
    if up is not None and c2.button("Save as current version", type="primary", width="stretch"):
        now = dt.datetime.now()
        db.add_notes(_notes_filename(now), up.getvalue(), now.isoformat(timespec="minutes"))
        st.success("Saved — this is now your current version.")
        st.rerun()

    vers = db.notes_versions()
    if not vers.empty:
        st.markdown("<hr class='page-rule'>", unsafe_allow_html=True)
        st.markdown("##### Version history")
        for _, v in vers.iterrows():
            cols = st.columns([0.6, 4.2, 2, 1.3, 1.1], vertical_alignment="center")
            cols[0].markdown(f"**v{int(v['version'])}**")
            cols[1].write(v["filename"])
            cols[2].write(f":gray[{str(v['uploaded_at']).replace('T', ' ')}]")
            _, vdata = db.get_notes(int(v["version"]))
            cols[3].download_button("Download", data=vdata, file_name=v["filename"],
                                    key=f"nv_{v['version']}", mime=DOCX_MIME, width="stretch")
            if cols[4].button("Delete", key=f"nvdel_{v['version']}", width="stretch"):
                db.delete_notes_version(int(v["version"]))
                st.rerun()


# ================================================================= RESOURCES
# Files live in the DB (not the git repo), so they travel to every device via
# Postgres and stay behind your password. Personal cross-device access only.
RESOURCES_DIR = os.environ.get("CFA_RESOURCES_DIR", os.path.expanduser("~/Desktop/CFA"))
RESOURCE_EXTS = (".pdf", ".xlsx", ".xlsm", ".csv", ".png", ".jpg", ".jpeg")
# Tidy up messy source filenames on the way into the app (download-facing names).
PRETTY_NAMES = {
    "CFA_L2_Study_Planner_2.xlsx": "CFA L2 Study Planner.xlsx",
    "aug 2025.png": "CFA Level I - Result (Aug 2025).png",
    "may 2026.png": "CFA Level I - Result (May 2026, Pass).png",
}
_MIMES = {".pdf": "application/pdf",
          ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
          ".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg"}


def _res_kind(name):
    """(group label, sort key, icon) for a resource, so the grid reads tidily."""
    n = name.lower()
    if n.endswith((".png", ".jpg", ".jpeg")):
        return ("Level I result", 3, "🎯")
    if n.endswith((".xlsx", ".xlsm", ".csv")):
        return ("Study planner", 2, "📊")
    if "quicksheet" in n:
        return ("Quicksheet", 1, "⚡")
    return ("Reading material", 0, "📕")


def page_resources():
    page_header("Reference", "Resources")
    if not require_unlock():
        return
    st.caption("Your materials, stored in the app database so they follow you to every device "
               "you deploy to — behind your password, for your personal use only.")
    have = db.list_resources()
    if not have.empty:
        have = have.assign(_sort=[_res_kind(n)[1] for n in have["name"]]) \
                   .sort_values(["_sort", "name"]).reset_index(drop=True)
        rows = list(have.iterrows())
        for i in range(0, len(rows), 2):
            cols = st.columns(2)
            for col, (_, r) in zip(cols, rows[i:i + 2]):
                grp, _s, icon = _res_kind(r["name"])
                ext = os.path.splitext(r["name"])[1].lower()
                with col:
                    st.markdown(
                        f"<div class='rescard'><span class='rc-name'>{icon}&nbsp; {r['name']}</span>"
                        f"<br><span class='rc-meta'>{grp} · {r['size'] / 1e6:.1f} MB</span></div>",
                        unsafe_allow_html=True)
                    st.download_button("Download", data=db.get_resource(r["name"]),
                                       file_name=r["name"],
                                       mime=_MIMES.get(ext, "application/octet-stream"),
                                       key=f"res_{r['name']}", width="stretch")
        st.markdown("<hr class='page-rule'>", unsafe_allow_html=True)

    if os.path.isdir(RESOURCES_DIR):
        st.caption(f"Local folder detected — `{RESOURCES_DIR}`")
        if st.button("Load / refresh files from my CFA folder into the app"):
            files = [f for f in glob.glob(os.path.join(RESOURCES_DIR, "**", "*"), recursive=True)
                     if f.lower().endswith(RESOURCE_EXTS)]
            now = dt.datetime.now().isoformat(timespec="minutes")
            for f in files:
                base = os.path.basename(f)
                with open(f, "rb") as fh:
                    db.upsert_resource(PRETTY_NAMES.get(base, base), fh.read(), now)
            st.success(f"Loaded {len(files)} file(s) into the app database.")
            st.rerun()
    elif have.empty:
        st.info("No materials loaded yet. Run the app locally (where your CFA folder lives), "
                "unlock, and use the loader to push them into the database once — then they're "
                "available wherever you deploy.")


# ----------------------------------------------------------------- dispatch
{"Today": page_today, "Calendar": page_calendar, "Curriculum": page_curriculum,
 "Reviews": page_reviews, "Drill Log": page_drill, "Mocks": page_mocks,
 "Analytics": page_analytics, "Notes": page_notes, "Resources": page_resources}[page]()
